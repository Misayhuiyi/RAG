"""Multi-format document loader.

Supports PDF (PyMuPDF), DOCX (python-docx), and Markdown files.
Returns unified LangChain Document objects.
"""

import logging
from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Load documents from PDF, DOCX, and Markdown files.

    Usage:
        loader = DocumentLoader()
        docs = loader.load_directory("./data/documents")
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt"}

    def __init__(self, supported_formats: Optional[list[str]] = None):
        self._supported = set(
            ext.lower() for ext in (supported_formats or self.SUPPORTED_EXTENSIONS)
        )

    # ── Public API ───────────────────────────────────────────────────

    def load_directory(self, dir_path: str | Path) -> List[Document]:
        """Recursively load all supported documents from a directory."""
        dir_path = Path(dir_path)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {dir_path}")

        all_docs: List[Document] = []
        for file_path in dir_path.rglob("*"):
            if file_path.is_file() and self._is_supported(file_path):
                try:
                    docs = self.load_file(file_path)
                    all_docs.extend(docs)
                    logger.info("Loaded %d chunks from %s", len(docs), file_path.name)
                except Exception as e:
                    logger.error("Failed to load %s: %s", file_path.name, e)

        logger.info("Loaded %d total documents from %s", len(all_docs), dir_path)
        return all_docs

    def load_file(self, file_path: str | Path) -> List[Document]:
        """Load a single document file. Dispatches to format-specific loader."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return self.load_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self.load_docx(file_path)
        elif ext in (".md", ".markdown"):
            return self.load_markdown(file_path)
        elif ext == ".txt":
            return self.load_text(file_path)
        else:
            raise ValueError(f"Unsupported format: {ext}")

    def load_pdf(self, file_path: str | Path) -> List[Document]:
        """Load a PDF file using PyMuPDF (fitz).

        Each page becomes a separate Document with page-number metadata.
        """
        import fitz  # PyMuPDF

        file_path = Path(file_path)
        docs: List[Document] = []

        with fitz.open(str(file_path)) as pdf:
            for page_idx in range(pdf.page_count):
                page = pdf[page_idx]
                text = page.get_text("text")
                if not text or not text.strip():
                    continue

                docs.append(Document(
                    page_content=text.strip(),
                    metadata={
                        "source": str(file_path.absolute()),
                        "file_name": file_path.name,
                        "format": "pdf",
                        "page": page_idx + 1,
                        "total_pages": pdf.page_count,
                    }
                ))

        return docs

    def load_docx(self, file_path: str | Path) -> List[Document]:
        """Load a Word document using python-docx.

        Paragraphs are grouped; tables are extracted with structure preserved.
        """
        from docx import Document as DocxDocument

        file_path = Path(file_path)
        docx = DocxDocument(str(file_path))

        paragraphs: list[str] = []
        # Extract paragraphs
        for para in docx.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract tables as structured text
        for table_idx, table in enumerate(docx.tables):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text.strip():
                paragraphs.append(table_text)

        full_text = "\n\n".join(paragraphs)
        if not full_text.strip():
            return []

        return [Document(
            page_content=full_text,
            metadata={
                "source": str(file_path.absolute()),
                "file_name": file_path.name,
                "format": "docx",
                "paragraphs_count": len(paragraphs),
            }
        )]

    def load_markdown(self, file_path: str | Path) -> List[Document]:
        """Load a Markdown file.

        Returns a single Document per file; chunking handles splitting later.
        """
        file_path = Path(file_path)
        text = file_path.read_text(encoding="utf-8")

        if not text.strip():
            return []

        return [Document(
            page_content=text,
            metadata={
                "source": str(file_path.absolute()),
                "file_name": file_path.name,
                "format": "markdown",
            }
        )]

    def load_text(self, file_path: str | Path) -> List[Document]:
        """Load a plain text file."""
        file_path = Path(file_path)
        text = file_path.read_text(encoding="utf-8")

        if not text.strip():
            return []

        return [Document(
            page_content=text,
            metadata={
                "source": str(file_path.absolute()),
                "file_name": file_path.name,
                "format": "text",
            }
        )]

    # ── Helpers ──────────────────────────────────────────────────────

    def _is_supported(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self._supported
